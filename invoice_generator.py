import csv
from docx import Document
from docx2pdf import convert


def load_clients(csv_path):
    clients = []

    with open(csv_path, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            quantite = int(row["quantite"])
            prix_unitaire = float(row["prix_unitaire"])
            row["montant_total"] = quantite * prix_unitaire
            clients.append(row)

    return clients


def create_invoice(client, invoice_number, output_folder="."):
    doc = Document()

    doc.add_heading("Facture", level=1)
    doc.add_paragraph(f"Numero de facture : {invoice_number}")
    doc.add_paragraph(f"Client : {client['nom']}")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"

    header_cells = table.rows[0].cells
    header_cells[0].text = "Produit"
    header_cells[1].text = "Quantite"
    header_cells[2].text = "Prix unitaire"
    header_cells[3].text = "Total"

    row_cells = table.add_row().cells
    row_cells[0].text = client["produit"]
    row_cells[1].text = client["quantite"]
    row_cells[2].text = f"{client['prix_unitaire']} EUR"
    row_cells[3].text = f"{client['montant_total']} EUR"

    docx_filename = f"{output_folder}/facture_{client['nom'].replace(' ', '_')}.docx"
    doc.save(docx_filename)
    print(f"Facture creee : {docx_filename}")

    convert(docx_filename)
    print(f"Facture convertie en PDF")


if __name__ == "__main__":
    clients = load_clients("clients.csv")
    for index, client in enumerate(clients, start=1):
        invoice_number = f"FACT-{index:03d}"
        create_invoice(client, invoice_number)